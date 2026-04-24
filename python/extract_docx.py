from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from lxml.etree import _Element

INPUT_DIR = Path("data/input/actions_progress")
OUTPUT_PATH = Path("public/data/actions_progress.json")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def accept_revisions_text(element: _Element) -> str:
    """Extract paragraph text with tracked changes accepted.

    Includes text from <w:ins> (insertions), excludes <w:del> (deletions).
    python-docx's .text property ignores both, giving incomplete results.
    """
    parts: list[str] = []
    for node in element.iter():
        if node.tag == f"{{{W_NS}}}delText":
            continue
        if node.tag == f"{{{W_NS}}}t":
            parts.append(node.text or "")
    return "".join(parts).strip()


def parse_title(text: str) -> tuple[int, str, int | None]:
    m = re.match(r"WP(\d+)\s*\|\s*(.+?)(?:\s*\((\d+)\s*words?\))?\s*$", text, re.I)
    if not m:
        raise ValueError(f"Cannot parse title: {text!r}")
    wp_number = int(m.group(1))
    wp_name = re.sub(r"\s*\(words?\)\s*$", "", m.group(2)).strip()
    word_count = int(m.group(3)) if m.group(3) else None
    return wp_number, wp_name, word_count


def parse_action_section(text: str) -> tuple[str, str]:
    """Split a list paragraph into (action_numbers, body).

    Handles patterns like:
      "Action 68: ...", "Actions 2225: ...", "4344: ...",
      "#1, 2, 3, 5, 6 and 7: ...", "#49, 58: ...", "Action 63 and 65: ..."
    """
    m = re.match(
        r"((?:Actions?\s+)?#?[\d,\s]+(?:and\s+\d+)?)\s*[:.]?\s*(.*)",
        text,
        re.S,
    )
    if m:
        return m.group(1).strip().rstrip(":. "), m.group(2).strip()
    return "", text.strip()


def parse_table(table) -> list[dict]:
    rows = []
    for ri, row in enumerate(table.rows):
        cells = [accept_revisions_text(cell._element) for cell in row.cells]
        if ri == 0:
            continue
        if not any(cells):
            continue
        rows.append(
            {
                "action": cells[0] if len(cells) > 0 else "",
                "pathwayToDecision": cells[1] if len(cells) > 1 else "",
                "intergovernmentalConsideration": cells[2] if len(cells) > 2 else "",
                "writtenProducts": cells[3] if len(cells) > 3 else "",
            }
        )
    return rows


def extract_document(filepath: Path) -> dict:
    doc = Document(filepath)

    sections: dict[str, str | None] = {
        "title": None,
        "objective": None,
        "how_we_get_there": None,
        "progress_to_date": None,
        "next_steps": None,
    }

    current_section = None
    section_paragraphs: dict[str, list[tuple[str, str]]] = {
        k: [] for k in sections if k != "title"
    }

    for para in doc.paragraphs:
        text = accept_revisions_text(para._element)
        style = para.style.name

        if style == "Heading 1":
            sections["title"] = text
            current_section = None
            continue

        if style == "Heading 2":
            upper = text.upper()
            if "OBJECTIVE" in upper:
                current_section = "objective"
            elif "HOW WE GET THERE" in upper:
                current_section = "how_we_get_there"
            elif "PROGRESS" in upper:
                current_section = "progress_to_date"
            elif "NEXT STEPS" in upper:
                current_section = "next_steps"
            else:
                pass
            continue

        if not text:
            continue

        # "NEXT STEPS AND DECISIONS" sometimes appears as Normal text after
        # PROGRESS TO DATE instead of as a Heading 2
        if current_section == "progress_to_date" and re.match(
            r"NEXT STEPS", text, re.I
        ):
            current_section = "next_steps"
            continue

        # The "QUICK SUMMARY" line signals the end of prose sections
        if re.match(r"QUICK SUMMARY", text, re.I):
            current_section = None
            continue

        if current_section and current_section in section_paragraphs:
            section_paragraphs[current_section].append((style, text))

    title_text = sections["title"] or ""
    wp_number, wp_name, word_count = parse_title(title_text)

    objective_parts = [
        t for _, t in section_paragraphs["objective"]
    ]
    objective = " ".join(objective_parts)

    how_items = []
    for style, text in section_paragraphs["how_we_get_there"]:
        if style == "List Paragraph":
            nums, body = parse_action_section(text)
            how_items.append({"actionNumbers": nums, "text": body})
        else:
            how_items.append({"actionNumbers": "", "text": text})

    progress_parts = [t for _, t in section_paragraphs["progress_to_date"]]
    progress = " ".join(progress_parts)

    next_items = []
    for style, text in section_paragraphs["next_steps"]:
        if style == "List Paragraph":
            nums, body = parse_action_section(text)
            next_items.append({"actionNumbers": nums, "text": body})
        else:
            next_items.append({"actionNumbers": "", "text": text})

    summary_table = parse_table(doc.tables[0]) if doc.tables else []

    return {
        "workPackageNumber": wp_number,
        "workPackageName": wp_name,
        "wordCount": word_count,
        "objective": objective,
        "howWeGetThere": how_items,
        "progressToDate": progress,
        "nextStepsAndDecisions": next_items,
        "summaryTable": summary_table,
        "sourceFile": filepath.name,
    }


def main():
    files = sorted(INPUT_DIR.glob("*.docx"))
    print(f"Found {len(files)} .docx files in {INPUT_DIR}")

    seen_wps: dict[int, str] = {}
    results = []

    for f in files:
        record = extract_document(f)
        wp = record["workPackageNumber"]

        # Skip duplicate WP files — keep the latest (last alphabetically)
        if wp in seen_wps:
            print(
                f"  WP{wp}: replacing {seen_wps[wp]} with {f.name} (newer revision)"
            )
            results = [r for r in results if r["workPackageNumber"] != wp]

        seen_wps[wp] = f.name
        results.append(record)

    results.sort(key=lambda r: r["workPackageNumber"])

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as out:
        json.dump(results, out, ensure_ascii=False, indent=2)

    print(f"\nExtracted {len(results)} work package progress reports")
    for r in results:
        actions_how = len(r["howWeGetThere"])
        actions_next = len(r["nextStepsAndDecisions"])
        table_rows = len(r["summaryTable"])
        print(
            f"  WP{r['workPackageNumber']:2d} | {r['workPackageName'][:45]:45s} "
            f"| how={actions_how} next={actions_next} table={table_rows}"
        )

    print(f"\nJSON written to {OUTPUT_PATH.resolve()}")


if __name__ == "__main__":
    main()
